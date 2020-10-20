import pandas as pd
import docx
import os
import PySimpleGUI as sg
from openpyxl import load_workbook
from openpyxl import Workbook



# GUI for user to input the template form, candidate docs (folder) and the job name
def retrieveDirPath():
    layout = [[sg.Text('Job Name')],
              [sg.InputText()],
              [sg.Text('Template Form:')],  
              [sg.Input(), sg.FileBrowse()],  
              [sg.Text('Form Folder')],
              [sg.Input(), sg.FolderBrowse()],
              [sg.OK(), sg.Cancel()]]  
    
    window = sg.Window('Equality & Diversity Reader').Layout(layout)  
    
    # The Event Loop  
    while True:  
        event, values = window.Read()  
        if event is None:  
            break
        if event == 'OK':  
            template_path = values[1]
            folder_path = values[2]
            job_name = values[0]
            break
    window.Close()    
    return template_path, folder_path, job_name

# Search Directory for docs 
def getFolderDocs(directory):
    folder = os.listdir(directory)
    for i in folder:
        if i[-4:] != 'docx': # Only docx files
            folder.pop(i)
    return folder

# Read template E&D form and return 
def readTemplate(template):
    template_tables = template.tables
    demographic_categories = {}
    i = 1

    for table in template_tables:
        data = []
    # Iterates over each cell and adds to list
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    data.append(paragraph.text)
    # Removes empty strings in list
        data = [x for x in data if x]
    # Adds dictionary key as the table number, values being a list of possible answers
        demographic_categories['Question ' + str(i)] = data
        i += 1
    # Creates list of all possible answers in the form
    template_answers = list(demographic_categories.values())
    template_answers = [inner for outer in template_answers for inner in outer]  
    return demographic_categories, template_answers

# Create empty Dataframe from template
def createDF(column_names):  
    # Returns list of tuples of all values (answers) with their associated key (question)
    cat_keys = [(i,x) for i in demographic_categories for x in demographic_categories[i]]

    # Creates multilevel dataframe with top level being the question number and the lower level being the possible answers for that question
    df = pd.DataFrame(columns=pd.MultiIndex.from_tuples(cat_keys, names=('Questions', 'Answers')))

    # Fills df with 0 to specify integer values
    df.loc[0, :] = 0
    return df

# Retrieve candidates answers from a filled E&D form
def retrieveAnswers(doc, template_answers):
    tables = doc.tables
    data = []
    # Extract answers
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    data.append(paragraph.text)
    # Remove emtpy cells
    data = [x for x in data if x]
    # Finds indices for filled cells that are not in answer cells
    tick_index = [index for index, value in enumerate(data) if value not in template_answers]
    # Map to change each element (x) of list to x-1
    demographic_index = list(map(lambda x: x-1, tick_index))
    # Returns a list of the answers given in doc
    demographics = [e for i, e in enumerate(data) if i in demographic_index]
    return demographics, demographic_index

# Insert candidates answers to the Dataframe
def insertAnswers(template, df, answers, answers_index, row_index):
    # Start at 'Question 1' = start at index 0 
    question_no = 1
    i = 0
    # Fill row with 0s
    df.loc[row_index, :] = 0
    while i in range(0, len(answers)):
    # Question No. corresponds to df
        question = "Question " + str(question_no)
    # If the first answer is in question 1, updates the df count by +1
        if answers[i] in list(df[question].columns):
            df.loc[row_index, (question, answers[i])] += 1
            print('Added to {} {}'.format(question, answers[i]))
            i += 1
            question_no += 1
        else:
            question_no += 1
    # Catches error if question_no is more than the number of tables
    # Removes string when two answers in the same string, only takes first answer
            if question_no > len(template.tables):
                answers.pop(i)
                question_no = i + 1
            continue
    return df

def sumTotal(df):
    df.loc['Total', :] = df.sum()
    return df

def removeBlanks(df):
    df = df.loc[:, (df != 0).any(axis=0)]
    return df

# Returns the list of questions from template
def findTemplateQuestions(template):
    paras = template.paragraphs
    questions = []
    # Append all questions to list including sub-questions (a)/(b)
    for para in paras:
        if len(para.text) > 0:
            if '(a)' in para.text:
                questions.append(para.text)
            elif '(b)' in para.text:
                questions.append(para.text)
            elif len(para.text) < 29:
                if '?' not in para.text: # Ensures no questions are added to list
                    questions.append(para.text)
    questions = questions[1:]
    
    # Remove headings for sub-questions
    idx_remove = []
    for i, question in enumerate(questions):
        if question[0:3] == '(a)':
            idx_remove.append(i-1)
            questions[i] = question[4:]
        elif question[0:3] == '(b)':
            questions[i] = question[4:]

    questions = [e for i, e in enumerate(questions) if i not in idx_remove]
    return questions


#####   
    

while True:
    # Open GUI and retrieve paths 
    template_path, folder_path, job_name = retrieveDirPath()
    print(template_path)
    try:
        if template_path[-4:] == 'docx':
            template = docx.Document(template_path)
        else:
            break
    except FileNotFoundError:
        print('File not Found 1')
        break
    
    # Read template and create df based on questions and answers
    demographic_categories, possible_answers = readTemplate(template)
    df = createDF(demographic_categories)   

    # Retrieve files from job folder
    try:
        folder = getFolderDocs(folder_path)
    except FileNotFoundError:
        print('File not Found 2')
        break
        
    for i, file in enumerate(folder):
        try:
            doc = docx.Document('{}\{}'.format(folder_path, file))
            answers, answers_idx = retrieveAnswers(doc, possible_answers)
            df = insertAnswers(template, df, answers, answers_idx, i)
        except:
            print('Error in file {}'.format(file))
            pass
    
    df = sumTotal(df)
    df = removeBlanks(df)
    
    # Return the questions from the template
    questions = findTemplateQuestions(template)
    
    # Create workbook
    book = Workbook()
    book.save('{}.xlsx'.format(job_name))
    print('Workbook created')
    book = load_workbook('{}.xlsx'.format(job_name))
    writer = pd.ExcelWriter("{}.xlsx".format(job_name), engine='openpyxl')
    writer.book = book
    writer.sheets = dict((ws.title, ws) for ws in book.worksheets)
    
    # Insert each question into workbook
    for i in range(0, len(df.columns.levels[0])):
        temp = df.loc['Total', 'Question '+str(i+1)].transpose()
    #    temp.rename(questions[i], inplace=True)
        temp.index.name = questions[i]
        sheet = questions[i][0:30]
        temp.to_excel(writer, sheet_name='Candidate Answers', startcol=i*3)
    
    writer.save()
    print('Job Completed.')
    break





























